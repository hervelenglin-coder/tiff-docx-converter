"""
TIFF to DOCX Converter Web Application
Converts multi-page TIFF files to formatted Word documents using Google Vision OCR
"""

# Monkey-patch pour gevent (doit être en premier!)
import os
if os.environ.get('ASYNC_MODE') == 'gevent':
    from gevent import monkey
    monkey.patch_all()

from flask import Flask, render_template, request, jsonify, send_file
from flask_socketio import SocketIO, emit
from werkzeug.utils import secure_filename
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer
import os
import uuid
import base64
import requests
import shutil
import threading
import re
import time

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'tiff-converter-secret-key-2024')
app.config['UPLOAD_FOLDER'] = os.environ.get('UPLOAD_FOLDER', 'uploads')
app.config['OUTPUT_FOLDER'] = os.environ.get('OUTPUT_FOLDER', 'output')
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024
app.config['GOOGLE_API_KEY'] = os.environ.get('GOOGLE_API_KEY', 'AIzaSyAl6_RXCta4-y1jD1TKDtChZrumsgzoRM4')

# Mode async: 'gevent' pour production, 'threading' pour développement
async_mode = os.environ.get('ASYNC_MODE', 'threading')
socketio = SocketIO(app, cors_allowed_origins="*", async_mode=async_mode)

ALLOWED_EXTENSIONS = {'tif', 'tiff'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def parse_page_ranges(page_string):
    """Parse une chaîne de pages comme '1-3,5,8-10' en liste [1,2,3,5,8,9,10]"""
    if not page_string or not page_string.strip():
        return set()

    pages = set()
    parts = page_string.replace(' ', '').split(',')

    for part in parts:
        if not part:
            continue
        if '-' in part:
            try:
                start, end = part.split('-', 1)
                start = int(start)
                end = int(end)
                if start <= end:
                    pages.update(range(start, end + 1))
            except ValueError:
                continue
        else:
            try:
                pages.add(int(part))
            except ValueError:
                continue

    return pages

def emit_progress(session_id, step, progress, message, status='processing'):
    """Émettre un événement de progression vers le client"""
    import sys
    data = {
        'session_id': session_id,
        'step': step,
        'progress': progress,
        'message': message,
        'status': status
    }
    # Émettre à la room spécifique ET en broadcast pour assurer la réception
    socketio.emit('progress_update', data, room=session_id)
    socketio.emit('progress_update', data, namespace='/')
    print(f"[PROGRESS] {session_id[:8]} - {step}: {progress}% - {message} ({status})", flush=True)
    sys.stdout.flush()

def convert_tiff_to_png(tiff_path, output_dir, session_id):
    png_files = []
    img = Image.open(tiff_path)
    n_frames = getattr(img, 'n_frames', 1)

    for i in range(n_frames):
        img.seek(i)
        output_path = os.path.join(output_dir, f'page_{i+1:03d}.png')
        if img.mode == '1':
            rgb_img = img.convert('RGB')
            rgb_img.save(output_path, 'PNG')
        else:
            img.save(output_path, 'PNG')
        png_files.append(output_path)
        progress = int((i + 1) / n_frames * 100)
        emit_progress(session_id, 'conversion', progress, f'Conversion page {i+1}/{n_frames}')

    return png_files

def google_vision_ocr_with_layout(image_path, api_key):
    """OCR avec détection de la mise en page, tableaux et scores de confiance"""
    with open(image_path, 'rb') as f:
        image_content = base64.b64encode(f.read()).decode('utf-8')

    url = f'https://vision.googleapis.com/v1/images:annotate?key={api_key}'
    payload = {
        'requests': [{
            'image': {'content': image_content},
            'features': [{'type': 'DOCUMENT_TEXT_DETECTION'}],
            # Option 1: Hint de langue française pour améliorer la reconnaissance
            'imageContext': {
                'languageHints': ['fr', 'en']
            }
        }]
    }

    response = requests.post(url, json=payload, timeout=120)
    result = response.json()

    # Debug: afficher le statut de la réponse
    print(f"[OCR DEBUG] HTTP Status: {response.status_code}")

    # Debug: afficher l'erreur si présente
    if 'error' in result:
        print(f"[OCR ERROR] API Error: {result['error']}")
        raise Exception(f"API Error: {result['error']}")

    # Vérifier les erreurs dans la réponse
    if 'responses' in result and result['responses']:
        resp = result['responses'][0]
        if 'error' in resp:
            print(f"[OCR ERROR] Response Error: {resp['error']}")
            raise Exception(f"Response Error: {resp['error']}")
        # Debug: afficher si du texte a été trouvé
        has_text = 'fullTextAnnotation' in resp
        print(f"[OCR DEBUG] Text found: {has_text}")
        if has_text:
            text_preview = resp['fullTextAnnotation']['text'][:100].replace('\n', ' ')
            print(f"[OCR DEBUG] Preview: {text_preview}...")
    else:
        print(f"[OCR DEBUG] No responses in result: {list(result.keys())}")

    data = {
        'full_text': '',
        'paragraphs': [],
        'blocks': [],  # Nouveau: blocs avec type (TABLE, TEXT, etc.)
        'width': 0,
        'height': 0,
        'detected_languages': [],
        'avg_confidence': 0
    }

    if 'responses' in result and result['responses']:
        resp = result['responses'][0]

        if 'fullTextAnnotation' in resp:
            data['full_text'] = resp['fullTextAnnotation']['text']

            pages = resp['fullTextAnnotation'].get('pages', [])
            if pages:
                page = pages[0]
                data['width'] = page.get('width', 1000)
                data['height'] = page.get('height', 1000)

                # Récupérer les langues détectées
                page_prop = page.get('property', {})
                detected_langs = page_prop.get('detectedLanguages', [])
                data['detected_languages'] = [
                    {'code': lang.get('languageCode', ''), 'confidence': lang.get('confidence', 0)}
                    for lang in detected_langs
                ]

                all_confidences = []

                for block in page.get('blocks', []):
                    # Option 3: Récupérer le type de bloc (TABLE, TEXT, PICTURE, etc.)
                    block_type = block.get('blockType', 'TEXT')
                    block_confidence = block.get('confidence', 0)
                    block_vertices = block.get('boundingBox', {}).get('vertices', [])

                    block_data = {
                        'type': block_type,
                        'confidence': block_confidence,
                        'paragraphs': []
                    }

                    if block_vertices and len(block_vertices) >= 4:
                        block_data['x'] = block_vertices[0].get('x', 0)
                        block_data['y'] = block_vertices[0].get('y', 0)
                        block_data['x2'] = block_vertices[2].get('x', 0)
                        block_data['y2'] = block_vertices[2].get('y', 0)

                    for para in block.get('paragraphs', []):
                        para_text = ""
                        para_confidence = para.get('confidence', 0)
                        vertices = para.get('boundingBox', {}).get('vertices', [])
                        word_details = []

                        for word in para.get('words', []):
                            word_text = ""
                            word_confidence = word.get('confidence', 0)
                            all_confidences.append(word_confidence)

                            for symbol in word.get('symbols', []):
                                word_text += symbol.get('text', '')

                            word_details.append({
                                'text': word_text,
                                'confidence': word_confidence
                            })
                            para_text += word_text + " "

                        if vertices and len(vertices) >= 4 and para_text.strip():
                            x = vertices[0].get('x', 0)
                            y = vertices[0].get('y', 0)
                            x2 = vertices[2].get('x', x)
                            y2 = vertices[2].get('y', y)

                            para_data = {
                                'text': para_text.strip(),
                                'x': x,
                                'y': y,
                                'width': x2 - x,
                                'height': y2 - y,
                                'x_percent': x / data['width'] * 100 if data['width'] > 0 else 0,
                                'y_percent': y / data['height'] * 100 if data['height'] > 0 else 0,
                                'confidence': para_confidence,
                                'words': word_details,
                                'block_type': block_type  # Associer le type de bloc au paragraphe
                            }

                            data['paragraphs'].append(para_data)
                            block_data['paragraphs'].append(para_data)

                    data['blocks'].append(block_data)

                # Calculer la confiance moyenne
                if all_confidences:
                    data['avg_confidence'] = sum(all_confidences) / len(all_confidences)

    # Trier par position Y puis X
    data['paragraphs'].sort(key=lambda p: (p['y'], p['x']))

    return data

def set_cell_border(cell, border_size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def detect_layout_zones(paragraphs, page_width):
    """Détecter les zones de mise en page (colonnes, en-têtes, tableaux)"""
    if not paragraphs:
        return []

    zones = []
    current_row = []
    last_y = -100
    row_threshold = 2.5  # % de hauteur pour considérer même ligne (plus précis)

    for para in paragraphs:
        y_pct = para['y_percent']

        # Nouvelle ligne si trop loin verticalement
        if abs(y_pct - last_y) > row_threshold and current_row:
            zones.append(current_row)
            current_row = []

        current_row.append(para)
        last_y = y_pct

    if current_row:
        zones.append(current_row)

    return zones


def is_table_zone(zone):
    """Déterminer si une zone est un tableau basé sur blockType de Google Vision"""
    for para in zone:
        if para.get('block_type') == 'TABLE':
            return True
    return False


def get_low_confidence_words(para_data, threshold=0.85):
    """Récupérer les mots avec une confiance faible"""
    low_conf_words = []
    for word in para_data.get('words', []):
        if word.get('confidence', 1.0) < threshold:
            low_conf_words.append(word['text'])
    return low_conf_words

def detect_table_structure(zones):
    """Détecter si plusieurs zones consécutives forment un tableau"""
    if len(zones) < 2:
        return []

    table_groups = []
    current_table = []

    for i, zone in enumerate(zones):
        # Si la zone a plusieurs colonnes, c'est potentiellement une ligne de tableau
        if len(zone) >= 2:
            current_table.append((i, zone))
        else:
            if len(current_table) >= 2:
                table_groups.append(current_table)
            current_table = []

    if len(current_table) >= 2:
        table_groups.append(current_table)

    return table_groups

def get_column_count(zones):
    """Déterminer le nombre de colonnes basé sur les zones"""
    if not zones:
        return 1
    col_counts = [len(z) for z in zones if len(z) > 1]
    if col_counts:
        return max(set(col_counts), key=col_counts.count)
    return 1

def create_formatted_document(ocr_data, page_num, total_pages, image_path=None, conversion_info=None):
    """Créer un document Word avec image + mise en page structurée style AMDEC"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Ajouter la page de résumé en première page du document
    if conversion_info:
        add_summary_page(doc, conversion_info)
        doc.add_page_break()

    # 1. Ajouter l'image originale
    if image_path and os.path.exists(image_path):
        # Réduire les marges pour l'image
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        section.top_margin = Cm(0.5)

        doc.add_picture(image_path, width=Cm(20))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Saut de page avant le texte formaté
        doc.add_page_break()

        # Restaurer les marges pour le texte
        # Note: on ne peut pas changer les marges après, donc on utilise l'indentation

    # Titre de la section texte avec indicateur de confiance
    avg_confidence = ocr_data.get('avg_confidence', 0)
    confidence_color = RGBColor(0, 128, 0) if avg_confidence >= 0.9 else RGBColor(200, 150, 0) if avg_confidence >= 0.75 else RGBColor(200, 0, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"══════════════════════════════════════════")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"TEXTE FORMATÉ - Page {page_num}/{total_pages}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Afficher la confiance moyenne si disponible
    if avg_confidence > 0:
        run = p.add_run(f"  [Confiance: {avg_confidence*100:.0f}%]")
        run.font.size = Pt(10)
        run.font.color.rgb = confidence_color

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"══════════════════════════════════════════")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    paragraphs = ocr_data.get('paragraphs', [])

    if not paragraphs:
        p = doc.add_paragraph()
        run = p.add_run("[Aucun texte détecté sur cette page]")
        run.italic = True
        run.font.size = Pt(11)
        return doc

    # Analyser les zones de mise en page
    zones = detect_layout_zones(paragraphs, ocr_data.get('width', 1000))

    # Variables pour suivre le contexte
    last_y_percent = 0
    in_header_section = True  # Les premières lignes sont souvent l'en-tête

    for zone_idx, row in enumerate(zones):
        if not row:
            continue

        # Trier par position X
        row.sort(key=lambda p: p['x'])

        first_para = row[0]
        y_pct = first_para['y_percent']

        # Détecter la fin de la section en-tête (après ~15% de la page)
        if y_pct > 15:
            in_header_section = False

        # Ajouter un espacement si grand saut vertical
        if y_pct - last_y_percent > 3 and zone_idx > 0:
            doc.add_paragraph()

        # Option 3: Détection améliorée des tableaux via blockType
        is_table = is_table_zone(row) or len(row) >= 2

        # Cas 1: Tableau détecté (par Google Vision ou multi-colonnes)
        if is_table:
            table = doc.add_table(rows=1, cols=len(row))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for col_idx, para in enumerate(row):
                cell = table.cell(0, col_idx)
                cell_para = cell.paragraphs[0]

                text = para['text']
                low_conf_words = get_low_confidence_words(para)

                # Alignement basé sur la position dans la ligne
                if col_idx == 0:
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif col_idx == len(row) - 1:
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Si le bloc est identifié comme TABLE par Google Vision, ajouter bordures
                if para.get('block_type') == 'TABLE' or in_header_section or is_header_text(text):
                    set_cell_border(cell, "4")
                    run = cell_para.add_run(text)
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    run = cell_para.add_run(text)
                    run.font.size = Pt(10)

                run.font.name = 'Arial'

                # Marquer visuellement les mots à faible confiance
                if low_conf_words:
                    run = cell_para.add_run(f" [?]")
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(200, 150, 0)

        # Cas 2: Un seul élément
        else:
            para = row[0]
            text = para['text']
            x_pct = para['x_percent']
            low_conf_words = get_low_confidence_words(para)

            p = doc.add_paragraph()

            # Déterminer le style basé sur le contenu et la position

            # Section title (souligné, comme "Commentaires générant :")
            if is_section_title(text):
                run = p.add_run(text)
                run.bold = True
                run.underline = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)

            # Titre principal (centré, gros)
            elif is_title_text(text):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Arial'
                p.paragraph_format.space_before = Pt(6)

            # Item numéroté (1., 2., Rep 03 :, etc.)
            elif is_numbered_item(text):
                # Légère indentation
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(text)
                run.font.size = Pt(11)
                run.font.name = 'Arial'

            # Ligne de continuation (texte indenté)
            elif x_pct > 15:
                # Indentation proportionnelle à la position X
                indent = Cm(max(1, (x_pct - 10) / 15))
                p.paragraph_format.left_indent = indent
                run = p.add_run(text)
                run.font.size = Pt(11)
                run.font.name = 'Arial'

            # En-tête (tableau du haut)
            elif in_header_section:
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Arial'

            # Texte normal
            else:
                run = p.add_run(text)
                run.font.size = Pt(11)
                run.font.name = 'Arial'

            # Indicateur visuel pour mots à faible confiance
            if low_conf_words and len(low_conf_words) <= 3:
                run = p.add_run(f"  [? {', '.join(low_conf_words)}]")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(200, 150, 0)
                run.italic = True

        last_y_percent = y_pct

    return doc

def is_section_title(text):
    """Détecter si le texte est un titre de section (souligné dans l'original)"""
    patterns = [
        r'^Commentaires?\s+(générant|spécifiques?)\s*:',
        r'^Observations?\s*:',
        r'^Remarques?\s*:',
        r'^Notes?\s*:',
        r'^Conclusions?\s*:',
        r'^Résumé\s*:',
    ]
    for pattern in patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True
    # Texte court finissant par ":"
    if len(text) < 40 and text.strip().endswith(':'):
        words = text.split()
        if len(words) <= 4:
            return True
    return False

def is_numbered_item(text):
    """Détecter si le texte est un item numéroté"""
    patterns = [
        r'^\d+\.\s',           # 1. , 2. , etc.
        r'^[a-z]\)\s',          # a) , b) , etc.
        r'^Rep\s+\d+\s*:',      # Rep 03 :
        r'^-\s',                # - item
        r'^•\s',                # • item
        r'^\(\d+\)',            # (1), (2), etc.
    ]
    for pattern in patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True
    return False

def is_reference_text(text):
    """Détecter si le texte est une référence/code"""
    patterns = [
        r'^\d{6}', r'^[A-Z]{2,4}-\d+', r'^\d+\.\d+\.\d+',
        r'^REV\s*\d', r'^IND\s*\d', r'^V\d+\.\d+'
    ]
    for pattern in patterns:
        if re.search(pattern, text):
            return True
    return False

def is_header_text(text):
    """Détecter si le texte est un en-tête de tableau/section"""
    patterns = [
        r'^REF\b', r'^N°', r'^PAGE\b', r'^INDICE\b', r'^REVISION\b',
        r'^DOCUMENT\b', r'^SYSTEME\b', r'^FOLIO\b', r'\bAMDEC\b',
        r'^DATE\b', r'^EMETTEUR\b', r'^VERIFICATEUR\b', r'^APPROBATEUR\b',
        r'^FONCTION\b', r'^MODE\b', r'^EFFET\b', r'^CAUSE\b', r'^GRAVITE\b',
        r'^FREQUENCE\b', r'^CRITICITE\b', r'^DETECTION\b', r'^MESURE\b',
        r'^REPERE\b', r'^DESIGNATION\b', r'^DESCRIPTION\b', r'^OBSERVATION\b',
        r'^EQUIPEMENT\b', r'^COMPOSANT\b', r'^DEFAILLANCE\b'
    ]
    text_upper = text.upper().strip()
    for pattern in patterns:
        if re.search(pattern, text_upper):
            return True
    # Tout en majuscules, court et pas un numéro simple
    if text.isupper() and 3 < len(text) < 40 and not text.isdigit():
        return True
    return False

def is_title_text(text):
    """Détecter si le texte est un titre principal"""
    if len(text) > 80:
        return False
    patterns = [
        r'ANALYSE DES MODES', r'GENERALITES', r'FONCTIONS REALISEES',
        r'ETUDE DE SECURITE', r'COMMENTAIRES', r'ANNEXE',
        r'INTRODUCTION', r'CONCLUSION', r'SOMMAIRE', r'TABLE DES MATIERES',
        r'CHAPITRE\s+\d', r'SECTION\s+\d', r'PARTIE\s+\d',
        r'TRANSMANCHE', r'EUROTUNNEL', r'TML',
        r'TOILETTES', r'ALIMENTATION', r'EAU'
    ]
    text_upper = text.upper()
    for pattern in patterns:
        if re.search(pattern, text_upper):
            return True
    return False

def add_summary_page(doc, conversion_info):
    """Ajouter une page de résumé de la conversion"""
    # Titre principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DOCUMENT CONVERTI")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Conversion TIFF vers DOCX avec OCR Google Vision")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tableau de résumé - ajout de lignes pour les stats OCR avancées
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Définir les largeurs
    table.columns[0].width = Cm(6)
    table.columns[1].width = Cm(10)

    # Ligne 1: Fichier source
    cell = table.cell(0, 0)
    cell.paragraphs[0].add_run("Fichier source").bold = True
    cell = table.cell(0, 1)
    cell.paragraphs[0].add_run(conversion_info.get('filename', '-'))

    # Ligne 2: Date de conversion
    cell = table.cell(1, 0)
    cell.paragraphs[0].add_run("Date de conversion").bold = True
    cell = table.cell(1, 1)
    cell.paragraphs[0].add_run(conversion_info.get('date', '-'))

    # Ligne 3: Pages totales
    cell = table.cell(2, 0)
    cell.paragraphs[0].add_run("Pages totales").bold = True
    cell = table.cell(2, 1)
    cell.paragraphs[0].add_run(str(conversion_info.get('total_pages', 0)))

    # Ligne 4: Pages avec OCR
    cell = table.cell(3, 0)
    cell.paragraphs[0].add_run("Pages avec OCR").bold = True
    cell = table.cell(3, 1)
    run = cell.paragraphs[0].add_run(str(conversion_info.get('ocr_pages', 0)))
    run.font.color.rgb = RGBColor(0, 128, 0)

    # Ligne 5: Pages exclues
    cell = table.cell(4, 0)
    cell.paragraphs[0].add_run("Pages exclues de l'OCR").bold = True
    cell = table.cell(4, 1)
    excluded = conversion_info.get('excluded_pages', [])
    if excluded:
        excluded_text = format_page_list_compact(excluded)
        run = cell.paragraphs[0].add_run(f"{len(excluded)} ({excluded_text})")
        run.font.color.rgb = RGBColor(200, 150, 0)
    else:
        cell.paragraphs[0].add_run("Aucune")

    # Ligne 6: Confiance moyenne OCR (nouvelle)
    cell = table.cell(5, 0)
    cell.paragraphs[0].add_run("Confiance OCR moyenne").bold = True
    cell = table.cell(5, 1)
    avg_conf = conversion_info.get('avg_confidence', 0)
    if avg_conf > 0:
        conf_text = f"{avg_conf*100:.1f}%"
        run = cell.paragraphs[0].add_run(conf_text)
        if avg_conf >= 0.9:
            run.font.color.rgb = RGBColor(0, 128, 0)  # Vert
        elif avg_conf >= 0.75:
            run.font.color.rgb = RGBColor(200, 150, 0)  # Orange
        else:
            run.font.color.rgb = RGBColor(200, 0, 0)  # Rouge
    else:
        cell.paragraphs[0].add_run("N/A")

    # Ligne 7: Langues détectées (nouvelle)
    cell = table.cell(6, 0)
    cell.paragraphs[0].add_run("Langues détectées").bold = True
    cell = table.cell(6, 1)
    languages = conversion_info.get('detected_languages', [])
    if languages:
        lang_texts = []
        for lang in languages[:3]:  # Max 3 langues
            code = lang.get('code', '').upper()
            conf = lang.get('confidence', 0)
            lang_texts.append(f"{code} ({conf*100:.0f}%)")
        cell.paragraphs[0].add_run(", ".join(lang_texts))
    else:
        cell.paragraphs[0].add_run("Non détectées")

    # Ligne 8: Options OCR utilisées (nouvelle)
    cell = table.cell(7, 0)
    cell.paragraphs[0].add_run("Options OCR").bold = True
    cell = table.cell(7, 1)
    options = ["Hint langue FR", "Scores confiance", "Détection tableaux"]
    run = cell.paragraphs[0].add_run(", ".join(options))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Appliquer le style aux cellules
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                for run in para.runs:
                    run.font.name = 'Arial'
                    if run.font.size is None:
                        run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Note explicative
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Structure du document :")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pour chaque page : Image originale → Texte formaté extrait par OCR")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(Les pages exclues contiennent uniquement l'image)")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 50)
    run.font.color.rgb = RGBColor(200, 200, 200)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Développé par Hervé Lenglin | Propulsé par Google Cloud Vision")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)


def format_page_list_compact(pages):
    """Formater une liste de pages de manière compacte (1,2,3,5 -> 1-3,5)"""
    if not pages:
        return ""
    if len(pages) == 1:
        return str(pages[0])

    pages = sorted(pages)
    ranges = []
    start = pages[0]
    end = pages[0]

    for i in range(1, len(pages)):
        if pages[i] == end + 1:
            end = pages[i]
        else:
            ranges.append(f"{start}" if start == end else f"{start}-{end}")
            start = pages[i]
            end = pages[i]

    ranges.append(f"{start}" if start == end else f"{start}-{end}")
    return ", ".join(ranges)


def create_simple_image_page(image_path):
    """Page avec image seule"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)

    doc.add_picture(image_path, width=Cm(20))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc

def create_image_plus_text_page(image_path, ocr_data, page_num, total_pages):
    """Page avec image puis texte sur page suivante"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)

    # Image
    doc.add_picture(image_path, width=Cm(20))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Saut de page
    doc.add_page_break()

    # Titre
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"══ TEXTE EXTRAIT - Page {page_num}/{total_pages} ══")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # Texte structuré
    full_text = ocr_data.get('full_text', '')
    if full_text:
        for line in full_text.split('\n'):
            if line.strip():
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(10)
                run.font.name = 'Arial'

    return doc

def merge_documents(doc_files, output_path):
    if not doc_files:
        return None
    master = Document(doc_files[0])
    composer = Composer(master)
    for doc_file in doc_files[1:]:
        doc = Document(doc_file)
        composer.append(doc)
    composer.save(output_path)
    return output_path

def process_tiff(tiff_path, api_key, session_id, output_mode='image_only', exclude_pages=None, original_filename=None):
    try:
        # Attendre que le client rejoigne la room
        socketio.sleep(0.5)

        work_dir = os.path.join(app.config['OUTPUT_FOLDER'], session_id)
        os.makedirs(work_dir, exist_ok=True)

        # Parse excluded pages
        excluded = parse_page_ranges(exclude_pages) if exclude_pages else set()

        # Conversion TIFF -> PNG
        emit_progress(session_id, 'conversion', 0, 'Conversion TIFF...', 'processing')
        socketio.sleep(0.1)  # Laisser le temps au message d'être envoyé
        png_files = convert_tiff_to_png(tiff_path, work_dir, session_id)
        total_pages = len(png_files)
        emit_progress(session_id, 'conversion', 100, f'{total_pages} pages', 'completed')

        # OCR
        ocr_results = []
        if output_mode == 'image_only':
            emit_progress(session_id, 'ocr', 100, 'OCR ignoré', 'completed')
            ocr_results = [{'full_text': '', 'paragraphs': []}] * total_pages
        else:
            pages_to_ocr = total_pages - len([p for p in excluded if 1 <= p <= total_pages])
            excluded_info = f" ({len(excluded)} exclues)" if excluded else ""
            emit_progress(session_id, 'ocr', 0, f'OCR Google Vision...{excluded_info}', 'processing')

            for i, png_file in enumerate(png_files):
                page_num = i + 1  # Pages start at 1
                progress = int((i + 1) / total_pages * 100)

                if page_num in excluded:
                    # Page exclue: pas d'OCR, juste un marqueur
                    ocr_results.append({'full_text': '', 'paragraphs': [], 'excluded': True})
                    emit_progress(session_id, 'ocr', progress, f'Page {page_num}/{total_pages} (exclue)', 'processing')
                    # Petit délai pour s'assurer que l'événement est envoyé
                    socketio.sleep(0.05)
                else:
                    try:
                        emit_progress(session_id, 'ocr', progress, f'OCR page {page_num}/{total_pages}...', 'processing')
                        result = google_vision_ocr_with_layout(png_file, api_key)
                        result['excluded'] = False
                        ocr_results.append(result)
                    except Exception as e:
                        ocr_results.append({'full_text': f'[Erreur: {str(e)}]', 'paragraphs': [], 'excluded': False})

            emit_progress(session_id, 'ocr', 100, 'OCR terminé', 'completed')
            socketio.sleep(0.1)  # S'assurer que le message "completed" est bien envoyé

        # Collecter les statistiques OCR avancées
        all_confidences = []
        all_languages = {}

        for ocr_data in ocr_results:
            if not ocr_data.get('excluded', False):
                # Collecter les confiances
                if ocr_data.get('avg_confidence', 0) > 0:
                    all_confidences.append(ocr_data['avg_confidence'])

                # Collecter les langues détectées
                for lang in ocr_data.get('detected_languages', []):
                    code = lang.get('code', '')
                    if code:
                        if code not in all_languages:
                            all_languages[code] = []
                        all_languages[code].append(lang.get('confidence', 0))

        # Calculer la confiance moyenne globale
        global_avg_confidence = sum(all_confidences) / len(all_confidences) if all_confidences else 0

        # Calculer la confiance moyenne par langue
        detected_languages = []
        for code, confidences in all_languages.items():
            avg_conf = sum(confidences) / len(confidences)
            detected_languages.append({'code': code, 'confidence': avg_conf})
        detected_languages.sort(key=lambda x: x['confidence'], reverse=True)

        # Préparer les métadonnées de conversion
        valid_excluded = sorted([p for p in excluded if 1 <= p <= total_pages])
        conversion_info = {
            'filename': original_filename or 'document.tif',
            'total_pages': total_pages,
            'excluded_pages': valid_excluded,
            'ocr_pages': total_pages - len(valid_excluded),
            'date': time.strftime('%d/%m/%Y %H:%M'),
            'avg_confidence': global_avg_confidence,
            'detected_languages': detected_languages
        }

        # Création documents
        emit_progress(session_id, 'formatting', 0, 'Mise en forme...', 'processing')
        doc_files = []

        for i, (png_file, ocr_data) in enumerate(zip(png_files, ocr_results)):
            page_num = i + 1
            is_excluded = ocr_data.get('excluded', False)

            if output_mode == 'image_only' or is_excluded:
                # Page exclue ou mode image seule: juste l'image
                doc = create_simple_image_page(png_file)
            elif output_mode == 'text_positioned':
                # Ajouter les infos de conversion uniquement sur la première page
                page_conversion_info = conversion_info if page_num == 1 else None
                doc = create_formatted_document(ocr_data, page_num, total_pages, png_file, page_conversion_info)
            else:
                doc = create_image_plus_text_page(png_file, ocr_data, page_num, total_pages)

            doc_path = os.path.join(work_dir, f'page_{i+1:03d}.docx')
            doc.save(doc_path)
            doc_files.append(doc_path)
            progress = int((i + 1) / total_pages * 100)
            status_text = f'Page {page_num}/{total_pages}'
            if is_excluded:
                status_text += ' (image seule)'
            emit_progress(session_id, 'formatting', progress, status_text)
            socketio.sleep(0.02)  # Petit délai pour l'UI

        emit_progress(session_id, 'formatting', 100, 'Terminé', 'completed')
        socketio.sleep(0.1)

        # Fusion
        emit_progress(session_id, 'merging', 0, 'Fusion...', 'processing')
        output_filename = f'document_converti_{session_id[:8]}.docx'
        output_path = os.path.join(work_dir, output_filename)
        merge_documents(doc_files, output_path)
        emit_progress(session_id, 'merging', 100, 'Fusion terminée', 'completed')
        socketio.sleep(0.1)

        emit_progress(session_id, 'complete', 100, f'{total_pages} pages converties', 'completed')
        socketio.sleep(0.1)
        return {'success': True, 'output_file': output_filename, 'total_pages': total_pages}

    except Exception as e:
        emit_progress(session_id, 'error', 0, f'Erreur: {str(e)}', 'error')
        return {'success': False, 'error': str(e)}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/preview', methods=['POST'])
def preview_file():
    """Génère un aperçu de toutes les pages et compte les pages"""
    if 'file' not in request.files:
        return jsonify({'error': 'Aucun fichier'}), 400

    file = request.files['file']

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Format non supporté'}), 400

    try:
        # Créer un dossier temporaire pour l'aperçu
        preview_id = str(uuid.uuid4())[:8]
        preview_dir = os.path.join(app.config['UPLOAD_FOLDER'], f'preview_{preview_id}')
        os.makedirs(preview_dir, exist_ok=True)

        # Sauvegarder temporairement le fichier
        filename = secure_filename(file.filename)
        temp_path = os.path.join(preview_dir, filename)
        file.save(temp_path)

        # Ouvrir le TIFF et compter les pages
        img = Image.open(temp_path)
        n_frames = getattr(img, 'n_frames', 1)

        # Générer les aperçus de toutes les pages
        previews = []
        max_width = 200  # Miniatures plus petites pour afficher plusieurs

        for i in range(n_frames):
            img.seek(i)

            # Convertir et redimensionner pour l'aperçu
            if img.mode == '1':
                preview_img = img.convert('RGB')
            else:
                preview_img = img.copy()

            # Redimensionner pour la miniature
            if preview_img.width > max_width:
                ratio = max_width / preview_img.width
                new_size = (max_width, int(preview_img.height * ratio))
                preview_img = preview_img.resize(new_size, Image.Resampling.LANCZOS)

            # Sauvegarder temporairement
            preview_path = os.path.join(preview_dir, f'preview_{i}.png')
            preview_img.save(preview_path, 'PNG')

            # Encoder en base64
            with open(preview_path, 'rb') as f:
                preview_base64 = base64.b64encode(f.read()).decode('utf-8')

            previews.append({
                'page': i + 1,
                'data': f'data:image/png;base64,{preview_base64}'
            })

        # Nettoyer les fichiers temporaires
        img.close()
        shutil.rmtree(preview_dir)

        return jsonify({
            'success': True,
            'total_pages': n_frames,
            'previews': previews
        })

    except Exception as e:
        # Nettoyer en cas d'erreur
        if 'preview_dir' in locals() and os.path.exists(preview_dir):
            shutil.rmtree(preview_dir)
        return jsonify({'error': str(e)}), 500

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'Aucun fichier'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'Aucun fichier sélectionné'}), 400

    if file and allowed_file(file.filename):
        session_id = str(uuid.uuid4())
        filename = secure_filename(file.filename)

        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], session_id)
        os.makedirs(upload_path, exist_ok=True)
        file_path = os.path.join(upload_path, filename)
        file.save(file_path)

        # Utiliser la clé API configurée et le mode text_positioned
        api_key = app.config['GOOGLE_API_KEY']
        output_mode = 'text_positioned'
        exclude_pages = request.form.get('exclude_pages', '')

        # Utiliser socketio.start_background_task pour compatibilité avec gevent
        socketio.start_background_task(process_tiff, file_path, api_key, session_id, output_mode, exclude_pages, filename)

        return jsonify({'success': True, 'session_id': session_id, 'filename': filename})

    return jsonify({'error': 'Format non supporté. Utilisez .tif ou .tiff'}), 400

@app.route('/download/<session_id>/<filename>')
def download_file(session_id, filename):
    file_path = os.path.join(app.config['OUTPUT_FOLDER'], session_id, filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return jsonify({'error': 'Fichier non trouvé'}), 404

@app.route('/cleanup/<session_id>', methods=['POST'])
def cleanup(session_id):
    try:
        for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
            path = os.path.join(folder, session_id)
            if os.path.exists(path):
                shutil.rmtree(path)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@socketio.on('connect')
def handle_connect():
    print('Client connected')

@socketio.on('join')
def handle_join(data):
    session_id = data.get('session_id')
    if session_id:
        from flask_socketio import join_room
        join_room(session_id)

@socketio.on('disconnect')
def handle_disconnect():
    print('Client disconnected')

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
    print("=" * 50)
    print("  TIFF to DOCX Converter")
    print("  http://localhost:5000")
    print("=" * 50)
    socketio.run(app, debug=True, host='0.0.0.0', port=5000)
